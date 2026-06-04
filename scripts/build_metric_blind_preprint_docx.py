#!/usr/bin/env python3
"""Build the Metric-Blind Leakage preprint DOCX from the Markdown draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "paper" / "metric_blind_leakage_preprint.md"
OUT = ROOT / "paper" / "metric_blind_leakage_preprint.docx"

PAGE_WIDTH_IN = 8.5
MARGIN_IN = 1.0
CONTENT_WIDTH_IN = PAGE_WIDTH_IN - 2 * MARGIN_IN


def main() -> int:
    doc = Document()
    configure_document(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    write_markdown(doc, lines)
    doc.save(OUT)
    print(OUT)
    return 0


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(10)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles.add_style("Figure Caption", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    equation = styles.add_style("Equation Block", 1)
    equation.font.name = "Consolas"
    equation.font.size = Pt(10)
    equation.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.text = "Metric-Blind Leakage preprint draft"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def write_markdown(doc: Document, lines: list[str]) -> None:
    idx = 0
    paragraph_buffer: list[str] = []
    while idx < len(lines):
        line = lines[idx].rstrip()

        if not line:
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            idx += 1
            continue

        if line.startswith("# "):
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            para = doc.add_paragraph(style="Title")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(para, line[2:].strip())
            idx += 1
            continue

        if line.startswith("## "):
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
            idx += 1
            continue

        if line.startswith("### "):
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            idx += 1
            continue

        if line.startswith("- "):
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.25)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.10
            add_inline_runs(para, line[2:].strip())
            idx += 1
            continue

        if line.startswith("|"):
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            table_lines = []
            while idx < len(lines) and lines[idx].startswith("|"):
                table_lines.append(lines[idx])
                idx += 1
            add_markdown_table(doc, table_lines)
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            add_image(doc, image_match.group(1), image_match.group(2))
            idx += 1
            continue

        if line == r"\[":
            flush_paragraph(doc, paragraph_buffer)
            paragraph_buffer = []
            eq_lines = []
            idx += 1
            while idx < len(lines) and lines[idx].rstrip() != r"\]":
                eq_lines.append(lines[idx].strip())
                idx += 1
            idx += 1
            add_equation(doc, " ".join(eq_lines))
            continue

        paragraph_buffer.append(line)
        idx += 1

    flush_paragraph(doc, paragraph_buffer)


def flush_paragraph(doc: Document, parts: list[str]) -> None:
    if not parts:
        return
    text = " ".join(part.strip() for part in parts).strip()
    if not text:
        return
    para = doc.add_paragraph()
    add_inline_runs(para, text)


def add_inline_runs(paragraph, text: str) -> None:
    chunks = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for chunk in chunks:
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(chunk)


def add_equation(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style="Equation Block")
    para.add_run(text)
    shade_paragraph(para, "F4F6F9")


def add_image(doc: Document, alt_text: str, rel_path: str) -> None:
    image_path = (SOURCE.parent / rel_path).resolve()
    if not image_path.exists():
        image_path = (SOURCE.parent / rel_path).resolve()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    width = Inches(5.9 if "decision_tree" in image_path.name else 5.5)
    run.add_picture(str(image_path), width=width)
    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run(alt_text)


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)

    widths = column_widths(len(rows[0]))
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_width(cell, widths[col_idx])
            if row_idx == 0:
                shade_cell(cell, "F2F4F7")
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            if col_idx >= 2:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(para, value)
            for run in para.runs:
                run.font.size = Pt(8.5 if len(rows[0]) > 5 else 9.5)
                if row_idx == 0:
                    run.bold = True
    doc.add_paragraph()


def column_widths(n_cols: int) -> list[float]:
    if n_cols == 8:
        return [1.20, 1.10, 0.56, 0.74, 0.58, 0.52, 1.15, 0.55]
    return [CONTENT_WIDTH_IN / n_cols for _ in range(n_cols)]


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9DEE6")


def set_cell_margins(table) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in {"top": "80", "bottom": "80", "start": "120", "end": "120"}.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), value)
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(para, fill: str) -> None:
    p_pr = para._p.get_or_add_pPr()
    shd = p_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def rgb(hex_color: str) -> RGBColor:
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )


if __name__ == "__main__":
    raise SystemExit(main())
